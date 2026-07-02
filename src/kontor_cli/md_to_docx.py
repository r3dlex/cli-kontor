"""Convert CS-intake markdown templates into clean .docx files.

Handles the subset of Markdown used by the Bug/Performance ticket templates:
YAML frontmatter, ATX headings, blockquotes, pipe tables, task-list checkboxes,
horizontal rules, and inline bold/italic/code/wikilinks (plus <br> in cells).

Source provenance: r3dlex/rib-workspace scripts/md_to_docx.py @ 3e3d082
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

_INLINE = re.compile(
    r"(\*\*.+?\*\*)"    # bold
    r"|(`[^`]+?`)"       # inline code
    r"|(\[\[.+?\]\])"   # wikilink
    r"|(_[^_]+?_)"       # italic (underscore)
)


def _run(paragraph: Any, text: str, bold: bool, italic: bool, code: bool) -> None:
    if not text:
        return
    r = paragraph.add_run(text)
    r.bold = bold or None
    r.italic = italic or None
    if code:
        r.font.name = "Consolas"


def add_runs(
    paragraph: Any,
    text: str,
    bold: bool = False,
    italic: bool = False,
    code: bool = False,
) -> None:
    """Append formatted runs for one logical line (no <br>), handling nesting."""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos : m.start()], bold, italic, code)
        b, c, wiki, ital = m.groups()
        if b is not None:
            add_runs(paragraph, b[2:-2], bold=True, italic=italic, code=code)
        elif c is not None:
            _run(paragraph, c[1:-1], bold, italic, True)
        elif wiki is not None:
            add_runs(paragraph, wiki[2:-2], bold=bold, italic=True, code=code)
        elif ital is not None:
            add_runs(paragraph, ital[1:-1], bold=bold, italic=True, code=code)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], bold, italic, code)


def fill_cell(cell: Any, text: str, *, header: bool = False) -> None:
    cell.text = ""
    parts = re.split(r"<br\s*/?>", text)
    for i, part in enumerate(parts):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        add_runs(para, part.strip())
        if header:
            for run in para.runs:
                run.bold = True
    if header:
        shade(cell, "D9E2F3")


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------

def shade(cell: Any, hex_color: str) -> None:
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def horizontal_rule(doc: Any) -> None:
    from docx.oxml import OxmlElement  # noqa: PLC0415
    from docx.oxml.ns import qn  # noqa: PLC0415

    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    borders.append(bottom)
    p_pr.append(borders)


# ---------------------------------------------------------------------------
# Block parsing
# ---------------------------------------------------------------------------

def strip_frontmatter(lines: list[str]) -> list[str]:
    if lines and lines[0].strip() == "---":
        for i in range(1, len(lines)):
            if lines[i].strip() == "---":
                return lines[i + 1 :]
    return lines


def is_table_row(line: str) -> bool:
    return line.lstrip().startswith("|")


def is_separator_row(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:|-]+\|?\s*", line)) and "-" in line


def split_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def emit_table(doc: Any, rows: list[str]) -> None:
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[2:]]
    ncols = len(header)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.autofit = True

    hrow = table.add_row().cells
    for i, htext in enumerate(header):
        if i < ncols:
            fill_cell(hrow[i], htext, header=True)
    for r in body:
        cells = table.add_row().cells
        for i in range(ncols):
            fill_cell(cells[i], r[i] if i < len(r) else "")
    doc.add_paragraph()


def emit_blockquote(doc: Any, qlines: list[str]) -> None:
    chunk: list[str] = []
    sentinel: list[str | None] = qlines + [None]
    for ql in sentinel:
        if ql is None or ql.strip() == "":
            if chunk:
                p = doc.add_paragraph(style="Quote")
                add_runs(p, " ".join(chunk))
                chunk = []
        else:
            chunk.append(ql)


def convert(md_path: str | Path, docx_path: str | Path) -> Path:
    """Convert a Markdown file to a .docx file.

    Returns the path of the written .docx file.
    """
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    lines = strip_frontmatter(
        Path(md_path).read_text(encoding="utf-8").splitlines()
    )
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    first_h1 = True
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped in ("---", "***", "___"):
            horizontal_rule(doc)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1 and first_h1:
                doc.add_heading(text, level=0)
                first_h1 = False
            elif level == 1:
                doc.add_heading(text, level=1)
            else:
                doc.add_heading(text, level=min(level - 1, 4))
            i += 1
            continue

        if stripped.startswith(">"):
            qlines: list[str] = []
            while i < n and lines[i].strip().startswith(">"):
                qlines.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            emit_blockquote(doc, qlines)
            continue

        if is_table_row(line) and i + 1 < n and is_separator_row(lines[i + 1]):
            rows: list[str] = []
            while i < n and is_table_row(lines[i]):
                rows.append(lines[i])
                i += 1
            emit_table(doc, rows)
            continue

        cb = re.match(r"^[-*]\s+\[( |x|X)\]\s+(.*)$", stripped)
        if cb:
            checked = cb.group(1).lower() == "x"
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("☒ " if checked else "☐ ")
            add_runs(p, cb.group(2))
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, bullet.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    out_path = Path(docx_path)
    doc.save(str(out_path))
    return out_path
